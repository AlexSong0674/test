# -*- coding: utf-8 -*-
"""Voyna (보이나) 기획서 v2.0 → DOCX 생성 스크립트"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\Users\송 하 준\Documents\알토대학원\수업과제\벤처 스타트업\기획서\Voyna_기획서_v2.docx"

# ── 헬퍼 함수 ─────────────────────────────────────────────────────────────

def set_font(run, name="맑은 고딕", size=None, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_spacing(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_cell_text(cell, text, bold=False, size=9.5, color=None, align=None):
    para = cell.paragraphs[0]
    para.clear()
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

def set_para_shading(para, hex_color):
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def add_para_border(para, color="CCCCCC"):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), color)
        pBdr.append(b)
    pPr.append(pBdr)

def add_hr(doc, color="0071E3"):
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)

def add_heading1(doc, text):
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 1"]
    run = para.add_run(text)
    set_font(run, size=15, bold=True, color=(0, 113, 227))
    para_spacing(para, before=14, after=6)
    return para

def add_heading2(doc, text):
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 2"]
    run = para.add_run(text)
    set_font(run, size=12, bold=True, color=(0, 81, 168))
    para_spacing(para, before=10, after=4)
    return para

def add_heading3(doc, text):
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 3"]
    run = para.add_run(text)
    set_font(run, size=10.5, bold=True, color=(68, 114, 196))
    para_spacing(para, before=8, after=3)
    return para

def parse_bold(text):
    import re
    result = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            result.append((text[last:m.start()], False))
        result.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        result.append((text[last:], False))
    if not result:
        result = [(text, False)]
    return result

def add_body(doc, text):
    para = doc.add_paragraph()
    for part_text, is_bold in parse_bold(text):
        run = para.add_run(part_text)
        set_font(run, size=10, bold=is_bold)
    para_spacing(para, before=0, after=4)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    for part_text, is_bold in parse_bold(text):
        run = para.add_run(part_text)
        set_font(run, size=10, bold=is_bold)
    para.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    para_spacing(para, before=0, after=2)
    return para

def add_numbered(doc, text, num):
    para = doc.add_paragraph()
    run_num = para.add_run(f"{num}. ")
    set_font(run_num, size=10, bold=False)
    for part_text, is_bold in parse_bold(text):
        run = para.add_run(part_text)
        set_font(run, size=10, bold=is_bold)
    para.paragraph_format.left_indent = Cm(0.5)
    para_spacing(para, before=0, after=2)
    return para

def add_quote(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "0071E3")
    pBdr.append(left)
    pPr.append(pBdr)
    set_para_shading(para, "EEF4FF")
    run = para.add_run(text)
    set_font(run, size=10, italic=True, color=(44, 62, 100))
    para_spacing(para, before=4, after=4)
    return para

def add_code_block(doc, text):
    lines = text.strip().split("\n")
    for line in lines:
        para = doc.add_paragraph()
        set_para_shading(para, "F4F4F4")
        add_para_border(para, "CCCCCC")
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(51, 51, 51)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Cm(0.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_checkbox_list(doc, items):
    for item in items:
        para = doc.add_paragraph()
        run = para.add_run("☐  " + item)
        set_font(run, size=10)
        para.paragraph_format.left_indent = Cm(0.5)
        para_spacing(para, before=0, after=2)

def make_table(doc, headers, rows, col_widths, header_bg="0071E3", row_alt_bg="EEF4FF"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, "AAAAAA")
        para = cell.paragraphs[0]
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(hdr)
        set_font(run, size=9.5, bold=True, color=(255, 255, 255))
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        cell.width = Cm(col_widths[i])

    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = row_alt_bg if r_idx % 2 == 0 else "FFFFFF"
        for i, (cell, val) in enumerate(zip(row.cells, row_data)):
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC")
            is_bold = val.startswith("**") and val.endswith("**")
            clean = val.strip("*")
            add_cell_text(cell, clean, bold=is_bold, size=9.5)
            cell.width = Cm(col_widths[i])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ── 문서 본문 ────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── 표지 ──
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Voyna (보이나)")
    set_font(title_run, size=30, bold=True, color=(0, 113, 227))
    title_para.paragraph_format.space_after = Pt(8)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("GPS 기반 여행 인증 배지 앱")
    set_font(sub_run, size=15, bold=False, color=(0, 81, 168))
    subtitle_para.paragraph_format.space_after = Pt(20)

    add_hr(doc)

    slogan_para = doc.add_paragraph()
    slogan_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = slogan_para.add_run('"발걸음이 기록이 되고, 기록이 추억이 된다"')
    set_font(s_run, size=12, italic=True, color=(80, 80, 80))
    slogan_para.paragraph_format.space_after = Pt(16)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        "초안: 2026-05-09  |  개정: 2026-05-12 (v2.0)  |  작성자: 창업자 (코딩 비전공 창업가)"
    )
    set_font(meta_run, size=9.5, color=(100, 100, 100), italic=True)
    meta_para.paragraph_format.space_after = Pt(8)

    nb_para = doc.add_paragraph()
    nb_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nb_run = nb_para.add_run(
        "Voyage + Navigation의 합성어, 한국어 '보이나'의 이중 의미"
    )
    set_font(nb_run, size=9, color=(120, 120, 120))

    doc.add_page_break()

    # ── 목차 ──
    add_heading1(doc, "목  차")
    toc = [
        "1. 서비스 개요",
        "2. 핵심 가치 제안",
        "3. 사용자 여정",
        "4. 기능 명세 (MVP)",
        "5. 앱 화면 구성",
        "6. 레벨 및 배지 시스템",
        "7. 수익 모델 로드맵",
        "8. 경쟁사 분석",
        "9. 기술 스택 검토",
        "10. 지도 라이브러리 비교",
        "11. 단계별 확장 전략",
        "12. 신규 아이디어 (V1.0+)",
        "13. 개발 일정 (로드맵)",
        "14. 개인정보 처리 방침 (요약)",
        "15. 리스크 및 대응 방안",
        "부록 A. MVP 개발 체크리스트",
        "부록 B. 주요 지표 (KPI)",
        "부록 C. 서울 초기 명소 50곳 (요약)",
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 1. 서비스 개요
    add_heading1(doc, "1. 서비스 개요")
    add_hr(doc)

    add_heading2(doc, "1.1 서비스 명")
    p = doc.add_paragraph()
    r1 = p.add_run("Voyna")
    set_font(r1, size=14, bold=True, color=(0, 113, 227))
    r2 = p.add_run(" (보이나)")
    set_font(r2, size=11, bold=False, color=(80, 80, 80))
    para_spacing(p, after=6)
    for item in [
        "어원: **Voyage** (항해/여정) + **Navigation** (탐험)",
        "한국어 '보이나' → 발견·인지의 의미 (예: \"추억이 보이나?\")",
        "글로벌·국내 모두 통하는 짧고 기억하기 쉬운 이름",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "1.2 한 줄 정의")
    add_quote(doc, "GPS 기반으로 여행지를 방문하면 자동으로 인증 배지를 획득하는 게이미피케이션 여행 앱")

    add_heading2(doc, "1.3 서비스 배경")
    for item in [
        "국내 여행 인구 증가 및 '인증' 문화 확산 (SNS 체크인, 스탬프 투어 유행)",
        "기존 종이 스탬프 투어의 디지털 전환 수요",
        "방문 기록을 게임처럼 수집하는 재미 요소에 대한 높은 수요",
        "외국인 관광객의 한국 여행 경험 디지털화 필요성",
        "대항해시대·포켓몬GO 등에서 검증된 '탐험+수집' 게임화 매커니즘의 일상 적용 가능성",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "1.4 타겟 사용자")
    make_table(doc,
        headers=["단계", "주요 타겟"],
        rows=[
            ["MVP", "국내 여행을 즐기는 20~40대 한국인"],
            ["확장 1단계", "국내 여행 마니아, 스탬프 투어 팬, 등산·트레킹 커뮤니티"],
            ["확장 2단계", "한국을 방문하는 외국인 관광객"],
            ["장기", "해외 여행지 도전을 원하는 한국인"],
        ],
        col_widths=[4, 12],
    )

    # 2. 핵심 가치 제안
    add_heading1(doc, "2. 핵심 가치 제안")
    add_hr(doc)

    add_heading2(doc, "2.1 슬로건")
    add_quote(doc, '"발걸음이 기록이 되고, 기록이 추억이 된다"')
    add_body(doc, "**서브 슬로건** (앱 스토어 / SNS 마케팅용)")
    for item in [
        '"걸음마다 추억이 보이나"',
        '"당신의 여정, Voyna에 담다"',
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "2.2 가치 키워드")
    for item in [
        "**여행자**: 내가 간 곳을 증명하고, 수집하는 재미",
        "**경쟁 요소**: 랭킹과 레벨로 다른 여행자와 겨루는 성취감",
        "**발견 요소**: 미수집 배지가 새로운 여행지로 유도",
        "**희소성**: 프리미어 배지·확인형 배지로 차별화된 소장 가치",
        "**추억**: 단순한 기록이 아니라 시간이 지나 다시 돌아볼 수 있는 디지털 여행 일기",
    ]:
        add_bullet(doc, item)

    # 3. 사용자 여정
    add_heading1(doc, "3. 사용자 여정")
    add_hr(doc)

    add_heading2(doc, "3.1 핵심 원칙")
    add_bullet(doc, "**앱 활성화 시에만 GPS 인식 및 배지 획득** 작동")
    add_bullet(doc, "사용자가 원할 때만 동작 → 배터리·프라이버시 부담 최소화", level=1)
    add_bullet(doc, "방해받고 싶지 않은 사용자도 수용", level=1)

    add_heading2(doc, "3.2 기본 플로우")
    add_code_block(doc,
        "앱 실행 → 홈 화면(레벨/배지 현황) → [여행 모드 ON]\n"
        "    → GPS 활성화 → 반경 내 여행지 감지 (50~200m 기준)\n"
        "    → [일반 배지] 자동 획득 애니메이션 → XP 적립\n"
        "    → [희귀 배지] '확인' 팝업 → 사용자 확인 → 특별 획득 연출\n"
        "    → 레벨/칭호 업데이트 확인\n"
        "    → [여행 모드 OFF]"
    )

    add_heading2(doc, "3.3 신규 사용자 온보딩")
    for i, item in enumerate([
        "회원가입 (카카오/구글 SNS 간편 로그인)",
        "닉네임 및 프로필 설정",
        "위치·알림 권한 요청 (이유 명확히 안내)",
        "주변 배지 획득 가능 장소 지도 미리 보기",
        "튜토리얼: 첫 배지 획득 유도 (가까운 명소 1곳)",
        "홈 화면으로 전환, 전국 배지 지도 확인",
    ], 1):
        add_numbered(doc, item, i)

    # 4. 기능 명세
    add_heading1(doc, "4. 기능 명세 (MVP)")
    add_hr(doc)

    add_heading2(doc, "4.1 MVP 범위 (우선 구현)")
    make_table(doc,
        headers=["기능", "설명", "우선순위"],
        rows=[
            ["회원가입/로그인", "카카오·구글 SNS 간편 로그인", "P0"],
            ["GPS 인식", "앱 활성화 시 위치 감지", "P0"],
            ["여행지 DB", "초기 주요 국내 명소 50~100곳", "P0"],
            ["**자동 배지 획득**", "일반 배지: 방문 인증 시 자동 부여", "P0"],
            ["**확인형 배지 획득**", "희귀·특별 배지: '확인' 버튼 팝업 후 획득", "P0"],
            ["배지 컬렉션", "획득/미획득 배지 목록 및 지도", "P0"],
            ["XP 및 레벨", "배지 획득 시 XP 적립, 레벨 표시", "P0"],
            ["홈 화면", "본인 레벨·칭호·최근 배지 요약", "P0"],
            ["이벤트 화면", "시즌·진행 중 이벤트 노출", "P1"],
            ["프로필", "레벨, 칭호, 배지 수, 방문 통계", "P1"],
            ["랭킹", "레벨/배지 수 기준 전체·지역 랭킹", "P1"],
            ["미수집 배지 추천", "근접 및 연계 기반 추천", "P1"],
            ["푸시 알림", "근처 미수집 배지 알림 (선택)", "P2"],
            ["소셜 공유", "배지 획득 시 SNS 공유 기능", "P2"],
            ["배너 광고", "비침습적 네이티브 광고 (Phase 1.5+)", "P2"],
        ],
        col_widths=[4, 9.5, 2.5],
    )

    add_heading2(doc, "4.2 배지 획득 방식 분기 (신규 사양)")
    add_code_block(doc,
        "[반경 내 진입 + GPS 정확도 충족]\n"
        "        │\n"
        "  ┌─────┴─────┐\n"
        "  ▼           ▼\n"
        "[일반 배지]    [희귀/특별 배지]\n"
        "자동 획득      확인 팝업\n"
        "+50 XP         '○○에 도착했어요! 이 순간을 기록할까요?'\n"
        "                       │\n"
        "                       ▼\n"
        "                [사용자 확인 버튼]\n"
        "                → 진동 + 사운드 + 풀스크린 애니메이션\n"
        "                → +150~500 XP"
    )
    add_body(doc, "**이유:**")
    for item in [
        "일반 배지의 자동성 → 빠르고 가벼운 UX",
        "희귀 배지의 '확인' → 특별한 순간임을 사용자가 인지 (소장 가치 강화)",
        "무분별한 우연 획득 방지 → 진짜 방문 의도가 있는 사용자에게 부여",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "4.3 MVP 이후 단계별 추가 기능")
    make_table(doc,
        headers=["단계", "기능"],
        rows=[
            ["V1.1", "칭호 시스템, 시즌 한정 배지, 친구 동반 보너스"],
            ["V1.2", "구독 모델, 프리미어 배지, 광고 제거"],
            ["V1.3", "여행 친구 팔로우, 방문 기록 타임라인"],
            ["V1.4", "테마 미션/퀘스트 시스템"],
            ["V2.0", "외국인 지원 (영어/일어/중국어), 해외 배지"],
            ["V2.1", "AR 배지 인증, 여행 일기 자동 생성"],
        ],
        col_widths=[3, 13],
    )

    doc.add_page_break()

    # 5. 앱 화면 구성
    add_heading1(doc, "5. 앱 화면 구성")
    add_hr(doc)
    add_body(doc, "MVP는 **하단 탭 4개** 구조 + 상단 알림/이벤트 영역으로 단순하게 시작한다.")
    add_code_block(doc,
        "┌──────────────────────────────────────────────┐\n"
        "│ ① 홈   ② 탐험 맵   ③ 배지   ④ 더보기            │\n"
        "│ (대시보드) (위치+주변) (컬렉션) (이벤트/설정)    │\n"
        "└──────────────────────────────────────────────┘"
    )

    add_heading2(doc, "① 홈 화면 (My Voyna)")
    for item in [
        "내 프로필: 닉네임, 레벨, 현재 칭호",
        "XP 진행 바: 다음 레벨까지 N XP",
        "최근 획득 배지 (3~5개 슬라이드)",
        "방문 통계: 누적 거리, 방문 지역 수",
        "다음 도전 카드: '이번 주 미션', '근처 배지 N개'",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "② 탐험 맵 (현재 위치 + 주변 배지)")
    for item in [
        "현 위치 중심 지도 (네이버맵 / 카카오맵)",
        "주변 미수집 배지 핀: 등급별 색상 구분, 미수집은 실루엣",
        "획득 배지 핀: 풀컬러 + 획득 일자",
        "여행 모드 ON/OFF 토글 (상단 우측)",
        "반경 필터: 1km / 3km / 10km / 전국",
        "카테고리 필터: 자연 / 역사 / 도시 / 음식",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "③ 배지 (컬렉션)")
    for item in [
        "종류별 탭: 전체 / 지역별 / 카테고리별 / 등급별",
        "배지 그리드: 획득(컬러) vs 미획득(실루엣)",
        "배지 상세 페이지: 디자인, 등급, 조건, XP, 획득자 통계, 관련 여행 정보",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "④ 더보기 (이벤트 + 설정)")
    for item in [
        "진행 중인 시즌 이벤트",
        "미션/퀘스트 목록 (V1.4+)",
        "친구·랭킹 (V1.3+)",
        "알림·언어·테마 설정",
        "위치 권한·개인정보 설정",
        "구독 관리 (V1.2+)",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # 6. 레벨 및 배지 시스템
    add_heading1(doc, "6. 레벨 및 배지 시스템")
    add_hr(doc)

    add_heading2(doc, "6.1 레벨 시스템")
    add_body(doc, "**레벨 범위**: 1 ~ 99")
    add_code_block(doc,
        "필요 XP(Lv n) = 100 × n^1.6   (소수점 버림)\n"
        "\n"
        "예시:\n"
        "  Lv  1 →  2:     100 XP\n"
        "  Lv  5 →  6:     952 XP\n"
        "  Lv 10 → 11:   2,512 XP\n"
        "  Lv 50 → 51:  39,821 XP\n"
        "  Lv 98 → 99: 146,779 XP"
    )

    add_heading2(doc, "6.2 배지 등급 및 획득 방식")
    make_table(doc,
        headers=["등급", "획득 방식", "XP", "시각 효과"],
        rows=[
            ["일반 (Common)", "자동 획득", "+50", "짧은 토스트 알림"],
            ["희귀 (Rare)", "'확인' 버튼 팝업", "+150", "진동 + 사운드 + 풀스크린 애니메이션"],
            ["특별 (Special)", "'확인' 팝업 (시즌·이벤트)", "+200~400", "컨페티 + 특수 사운드"],
            ["프리미어 (Premier)", "'확인' 팝업 (구독자 전용)", "+500", "고급 황금 애니메이션"],
            ["시즌 한정", "기간 내만 획득 가능", "+200~400", "시즌별 테마 연출"],
        ],
        col_widths=[3.5, 5, 2.5, 5],
    )

    add_heading2(doc, "6.3 칭호(타이틀) 시스템")
    make_table(doc,
        headers=["칭호", "부여 조건", "설명"],
        rows=[
            ["동네 탐험가", "배지 5개 달성", "시작하는 여행자"],
            ["국내 여행러", "3개 광역시·도 배지 보유", "지역 확장 중"],
            ["지도 마니아", "배지 30개 달성", "수집에 빠진 사람"],
            ["전국 방랑자", "모든 광역시·도 1개 이상", "전국 답파"],
            ["산악인", "국립공원 배지 10개 이상", "자연을 사랑하는 여행자"],
            ["도시 헌터", "5대 광역시 배지 완전 수집", "도시 여행 전문가"],
            ["고수 여행자", "Lv 50 달성", "중급 여행자"],
            ["전설의 여행자", "Lv 99 + 배지 200개 이상", "최고 등급 칭호"],
            ["시즌 챔피언", "시즌 랭킹 1위", "기간 한정 칭호"],
            ["선구자", "MVP 베타 가입자", "명예 칭호"],
        ],
        col_widths=[4, 7, 5],
    )
    add_quote(doc, "칭호는 프로필에 1개만 장착 가능 → 희소성과 자랑 욕구 자극")

    doc.add_page_break()

    # 7. 수익 모델 로드맵
    add_heading1(doc, "7. 수익 모델 로드맵")
    add_hr(doc)

    add_heading2(doc, "Phase 1 — 사용자 기반 확보 (0~4개월)")
    for item in [
        "**완전 무료 서비스**",
        "목표: MAU 3,000명, 배지 수집 데이터 축적",
        "수익: 없음 (투자 또는 자비)",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "Phase 1.5 — 광고 수익 시범 도입 (4~6개월)")
    for item in [
        "**비침습적 네이티브 광고** (배지 카드 사이, 더보기 탭 하단)",
        "여행 관련 타겟 광고 우선: 숙소·항공·렌터카·관광 상품",
        "예상 수익: MAU 1만 기준 월 50~150만 원",
        "광고 네트워크: AdMob, 카카오모먼트, 네이버 검색광고",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "Phase 2 — 구독 모델 도입 (6~12개월)")
    add_bullet(doc, "**Voyna+ 프리미엄 구독** (월 2,900원 / 연 24,900원)")
    for item in [
        "프리미어 배지 획득 권한",
        "**광고 완전 제거**",
        "고급 통계 (방문 히스토리, 이동 거리, 인사이트)",
        "신규 배지 사전 오픈 혜택",
        "친구 동반 보너스 XP 배율 증가",
    ]:
        add_bullet(doc, item, level=1)

    add_heading2(doc, "Phase 3 — 굿즈 및 B2B 협업 (12~24개월)")
    for item in [
        "**한정판 실물 굿즈**: 스티커팩, 엽서, 아크릴 키링",
        "**디지털 배지 컬렉션**: 특별 이벤트 배지 유료 판매",
        "**B2B 가맹 배지 시스템**: 지역 카페·맛집이 자체 배지 등록 (디지털 스탬프 카드)",
        "**지역 관광청·브랜드 협업 배지**",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "Phase 4 — 여행 콘텐츠 플랫폼 (24개월~)")
    for item in [
        "여행 코스 추천 (미수집 배지를 연결한 최적 동선)",
        "여행 크리에이터 콘텐츠 연동",
        "제휴 숙소·식당 예약 연결 (커미션 수익)",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "Phase 5 — 여행테크 글로벌 확장 (장기)")
    for item in [
        "외국인 대상 한국 여행 가이드 앱",
        "해외 여행지 배지 확대 (일본, 동남아, 유럽)",
        "여행 데이터 기반 광고·인사이트 서비스",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # 8. 경쟁사 분석 (신규)
    add_heading1(doc, "8. 경쟁사 분석")
    add_hr(doc)

    add_heading2(doc, "8.1 직접 경쟁사 (없음 — 블루오션)")
    add_quote(doc, "동일한 컨셉(GPS+자동 인증+배지 게이미피케이션+여행 전반)의 한국 시장 직접 경쟁사는 현재 부재.")

    add_heading2(doc, "8.2 간접 경쟁사 / 영감 사례")
    make_table(doc,
        headers=["서비스", "카테고리", "유사점", "Voyna 우위"],
        rows=[
            ["포켓몬GO", "AR 게임", "GPS+수집+게임화", "게임 전용 · 여행 본질 아님"],
            ["트랭글(Tranggle)", "등산 GPS", "정상 인증 배지", "등산 한정 · 일반 여행지 미지원"],
            ["대항해시대 (게임)", "탐험 RPG", "장소 발견·기록", "가상 세계 · 실제 여행 분리"],
            ["Foursquare / Swarm", "위치 체크인", "장소 체크인", "한국 정착 실패 · UX 노후"],
            ["여행레터, 트리플, 마이리얼트립", "여행 정보", "한국 여행 컨텐츠", "정보 중심 · 게이미피케이션 약함"],
            ["티켓플레이스 도장투어", "디지털 스탬프", "도장 수집", "단발성 이벤트 위주 · 지속성 약함"],
        ],
        col_widths=[3.5, 3, 4.5, 4.5],
    )

    add_heading2(doc, "8.3 차별화 포인트")
    for i, item in enumerate([
        "**여행 전체를 게임화** (등산/특정 카테고리 한정 아님)",
        "**자동+확인 이원화 배지 시스템** (게임적 재미 + 소장 가치)",
        "**블루오션 시장**: 직접 경쟁자 부재",
        "**B2B 확장성**: 관광청·지역 상권 협업",
        "**글로벌 잠재력**: 외국인 한국 여행 인바운드 자연 흡수",
    ], 1):
        add_numbered(doc, item, i)

    doc.add_page_break()

    # 9. 기술 스택
    add_heading1(doc, "9. 기술 스택 검토")
    add_hr(doc)

    add_heading2(doc, "9.1 MVP 빌드 도구 비교")
    make_table(doc,
        headers=["도구", "장점", "단점", "추천 용도"],
        rows=[
            ["Lovable", "AI 코드 생성, 빠른 프로토타이핑", "모바일 네이티브 GPS 제한", "관리자 대시보드"],
            ["Replit", "즉시 실행, 백엔드 API 구축", "UI/UX 완성도 낮음", "백엔드 API 서버"],
            ["FlutterFlow ⭐", "네이티브 앱 동시 생성, GPS 직접 지원", "러닝커브 존재", "MVP 앱 가장 적합"],
            ["Bubble", "강력한 노코드 웹앱", "모바일 변환 품질 낮음", "웹 버전 MVP"],
            ["Glide", "스프레드시트 기반 초고속", "GPS 활용 어려움", "부적합"],
        ],
        col_widths=[3.5, 5, 4.5, 3.5],
    )

    add_heading2(doc, "9.2 추천 스택")
    add_code_block(doc,
        "[Voyna MVP 추천 스택]\n"
        "- 앱 프론트엔드:  FlutterFlow (GPS 네이티브 + iOS/Android 동시 빌드)\n"
        "- 백엔드/DB:      Supabase (PostgreSQL)\n"
        "- 인증:           Supabase Auth (카카오·구글 소셜 로그인)\n"
        "- 지도(MVP):      카카오맵 API ── 최우선 (FlutterFlow 연동 우선)\n"
        "                  네이버맵 API ── V1.1 옵션 추가 (사용자 친숙도)\n"
        "                  구글맵 API ── 해외 확장 시\n"
        "- 광고 SDK:       AdMob (Phase 1.5+)\n"
        "- 분석:           Mixpanel 또는 Amplitude"
    )

    add_heading2(doc, "9.3 데이터베이스 설계 (핵심 테이블)")
    add_code_block(doc,
        "users           - id, nickname, level, xp, title, created_at\n"
        "locations       - id, name, lat, lng, radius, category, region\n"
        "badges          - id, location_id, grade, xp_reward, requires_confirmation, is_premium\n"
        "user_badges     - user_id, badge_id, obtained_at, lat, lng, photo_url\n"
        "xp_log          - user_id, amount, reason, created_at\n"
        "seasons         - id, name, start_at, end_at, theme\n"
        "subscriptions   - user_id, plan, expires_at, status\n"
        "events          - id, title, description, start_at, end_at, banner_url\n"
        "missions        - id, title, condition_json, reward_xp, badge_id\n"
        "friends         - user_id, friend_id, status, created_at"
    )
    add_body(doc, "**v2.0 주요 변경점**: `badges.requires_confirmation` 추가 (자동/확인 분기), `user_badges.photo_url` 추가, `events`/`missions` 신설")

    # 10. 지도 라이브러리 비교
    add_heading1(doc, "10. 지도 라이브러리 비교")
    add_hr(doc)

    make_table(doc,
        headers=["항목", "네이버맵", "카카오맵 ⭐", "구글맵"],
        rows=[
            ["국내 지도 품질", "★★★★★", "★★★★★", "★★★★☆"],
            ["국내 POI 데이터", "매우 풍부", "매우 풍부", "다소 부족"],
            ["사용자 친숙도", "★★★★★", "★★★★☆", "★★★☆"],
            ["무료 사용량", "월 200만 건", "월 300만 건", "월 $200 크레딧"],
            ["FlutterFlow 연동", "❌ 공식 지원 약함", "⚠️ 비공식 패키지", "✅ 공식 지원"],
            ["해외 지도 지원", "미지원", "미지원", "★★★★★"],
            ["MVP 개발 속도", "느림", "보통", "빠름"],
        ],
        col_widths=[4.5, 4, 4, 4],
    )

    add_heading2(doc, "추천 전략")
    add_code_block(doc,
        "MVP 1단계:    카카오맵 API\n"
        "              → FlutterFlow 연동 용이성 + 국내 POI 풍부 + 무료 한도 가장 넉넉\n"
        "\n"
        "V1.1:         네이버맵 옵션 추가 (사용자 선택형)\n"
        "              → 한국 사용자 친숙도가 가장 높음\n"
        "\n"
        "해외 확장 시: 구글맵 API로 분기 처리\n"
        "              → 국내는 카카오/네이버, 해외는 구글"
    )
    add_quote(doc, "사용자 친숙도만 본다면 네이버맵이 최선이지만, FlutterFlow 노코드 환경에서 네이버맵 연동은 커스텀 Dart 코드 작업이 필요합니다. MVP 출시 속도와 비전공 창업가의 개발 부담을 고려하면 카카오맵으로 시작 후 네이버맵 옵션 추가가 가장 현실적입니다.")

    doc.add_page_break()

    # 11. 단계별 확장 전략
    add_heading1(doc, "11. 단계별 확장 전략")
    add_hr(doc)

    add_heading2(doc, "11.1 국내 콘텐츠 확장")
    add_code_block(doc,
        "MVP 런칭:  주요 관광지 50~100곳 (서울 50곳 우선)\n"
        "V1.1:      전국 시군구 대표 명소 확장 (300곳+)\n"
        "V1.2:      지역 숨은 명소, 로컬 맛집·카페 연동\n"
        "V1.3:      지역 관광청 협업 공식 배지 코스\n"
        "V1.4:      사용자 큐레이션 (UGC) — 사용자가 새 장소 제안"
    )

    add_heading2(doc, "11.2 외국인 관광객 확장 (인바운드)")
    for item in [
        "다국어 지원: 영어, 일본어, 중국어(간체/번체)",
        "외국인 인기 명소 중심 배지 재구성",
        "한국 관광공사·지자체 협업 공식 파트너 배지",
        "입국 직후 공항에서 첫 배지 획득 유도 (인천공항 배지)",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "11.3 해외 여행지 확장 (아웃바운드)")
    for item in [
        "한국인 인기 국가부터: 일본 → 동남아 → 유럽",
        "현지 파트너사·로컬 크리에이터와 배지 큐레이션",
        "글로벌 여행자 커뮤니티로 성장",
    ]:
        add_bullet(doc, item)

    # 12. 신규 아이디어 (신규)
    add_heading1(doc, "12. 신규 아이디어 (V1.0+)")
    add_hr(doc)

    add_heading2(doc, "12.1 친구 동반 보너스 (V1.1)")
    add_bullet(doc, "친구와 함께 동일 장소 배지 획득 시 **양쪽 모두 추가 XP +30%**")
    add_bullet(doc, "4명 이상 함께 방문 시 한정 '단체 여행자' 배지")
    add_bullet(doc, "친구 초대 시 양쪽에 보상 (바이럴 성장)")

    add_heading2(doc, "12.2 테마 미션/퀘스트 (V1.4)")
    for item in [
        "시즌별/테마별 도전 과제 ('봄꽃 명소 5곳', '겨울 산 정상 3곳')",
        "미션 완료 트랙: 시즌 패스 형태",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "12.3 여행 일기 자동 생성 (V2.1)")
    add_bullet(doc, "획득 배지 + 사용자 사진 + 자동 메모 → 타임라인 일기")
    add_bullet(doc, "AI가 동선·시간 기반으로 짧은 여행기 초안 작성")
    add_bullet(doc, "SNS·블로그로 한 번에 공유")

    add_heading2(doc, "12.4 AR 배지 인증 (V2.1)")
    add_bullet(doc, "카메라를 켜면 **AR로 배지 마커가 떠 있음**")
    add_bullet(doc, "포켓몬GO 스타일 AR 효과 + 사진 촬영")
    add_bullet(doc, "일반 GPS 인증 외 보조 옵션")

    add_heading2(doc, "12.5 로컬 큐레이터 / UGC (V1.4)")
    add_bullet(doc, "사용자가 직접 새 장소 제안 → 채택 시 보상")
    add_bullet(doc, "채택자에게 '큐레이터' 칭호 + XP + 첫 획득자 명예")

    add_heading2(doc, "12.6 B2B 가맹 배지 (Phase 3)")
    add_bullet(doc, "지역 카페·맛집이 자체 배지 등록 (월 구독료 또는 건당 과금)")
    add_bullet(doc, "디지털 스탬프 카드 효과 + Voyna 콘텐츠 풍부화")

    add_heading2(doc, "12.7 사진 첨부 인증 (V2)")
    add_bullet(doc, "GPS 정확도가 떨어지는 실내·도심에서 사진 보조 인증")
    add_bullet(doc, "AI(Vision API)가 사진의 랜드마크 확인")

    doc.add_page_break()

    # 13. 개발 일정 (신규)
    add_heading1(doc, "13. 개발 일정 (로드맵)")
    add_hr(doc)

    add_heading2(doc, "13.1 MVP 개발 일정 (Gantt)")
    add_code_block(doc,
        "                    M1  M2  M3  M4  M5  M6\n"
        "기획·디자인 확정    ██▓\n"
        "DB 스키마 / 백엔드     ██▓\n"
        "FlutterFlow 앱 개발       ██████▓\n"
        "배지·장소 데이터 입력         ████\n"
        "GPS·인증 로직                    ███\n"
        "관리자 대시보드                      ██▓\n"
        "내부 테스트                              ██\n"
        "베타 (TestFlight)                          ██\n"
        "정식 출시                                    ▓"
    )

    add_heading2(doc, "13.2 단계별 마일스톤")
    make_table(doc,
        headers=["마일스톤", "시점", "핵심 산출물"],
        rows=[
            ["M1: 기획 완료", "0주", "기획서·와이어프레임·DB 스키마 확정"],
            ["M2: 백엔드 가동", "6주", "Supabase 가동, 명소 50곳 입력, API 동작"],
            ["M3: 앱 동작 (알파)", "12주", "로그인·지도·배지 획득 동작"],
            ["M4: 내부 테스트", "18주", "팀 내부·지인 10명 베타"],
            ["M5: 베타 출시", "22주", "TestFlight 100명 (선구자 칭호)"],
            ["M6: 정식 출시", "24주", "App Store·Play Store 출시"],
            ["M7: 광고 도입", "28주", "Phase 1.5 광고 도입"],
            ["M8: 구독 도입", "32주", "Phase 2 Voyna+ 출시"],
        ],
        col_widths=[4, 2.5, 9.5],
    )

    add_heading2(doc, "13.3 인적 자원 추정")
    make_table(doc,
        headers=["역할", "인원", "주요 업무"],
        rows=[
            ["PM / 창업자", "1", "기획·BM·운영·파트너십"],
            ["앱 개발자 (FlutterFlow)", "1~2", "앱 UI 및 GPS·배지 로직"],
            ["백엔드 개발자", "1 (외주 또는 풀스택)", "Supabase·관리자 대시보드"],
            ["디자이너", "1 (외주 또는 파트타임)", "배지 디자인·앱 UI"],
            ["콘텐츠 매니저", "1 (파트타임)", "명소 데이터·이벤트 기획"],
        ],
        col_widths=[4, 4, 8],
    )

    doc.add_page_break()

    # 14. 개인정보 처리 방침 (신규)
    add_heading1(doc, "14. 개인정보 처리 방침 (요약)")
    add_hr(doc)
    add_quote(doc, "출시 전 변호사 또는 개인정보보호 전문가 검토 필수. 아래는 정책 초안 골격.")

    add_heading2(doc, "14.1 수집 항목")
    make_table(doc,
        headers=["항목", "수집 시점", "목적", "보관 기간"],
        rows=[
            ["이메일·이름", "가입 시 (SNS 연동)", "계정 식별", "회원 탈퇴 시까지"],
            ["닉네임", "가입 시", "서비스 내 표시", "회원 탈퇴 시까지"],
            ["위치 정보 (위경도)", "여행 모드 ON 시", "배지 자동 인증", "30일 (인증 후 익명화)"],
            ["배지 획득 기록", "배지 획득 시", "컬렉션·통계 제공", "회원 탈퇴 시까지"],
            ["디바이스 정보", "앱 사용 시", "서비스 품질·디버깅", "90일"],
        ],
        col_widths=[3.5, 4, 4, 4.5],
    )

    add_heading2(doc, "14.2 핵심 원칙")
    for item in [
        "**최소 수집**: 서비스에 꼭 필요한 정보만",
        "**명시적 동의**: 위치·알림 권한은 사용 직전 안내 후 요청",
        "**명확한 OFF 옵션**: 여행 모드 OFF 시 위치 추적 즉시 중단",
        "**암호화 저장**: 모든 PII는 Supabase RLS + 암호화",
        "**광고 식별자 분리**: 마케팅 ID는 별도 옵트인 (Phase 1.5+)",
    ]:
        add_bullet(doc, item)

    add_heading2(doc, "14.3 위치 정보 특별 규정")
    for item in [
        "위치정보사업자 신고 필요 여부 검토 (개인위치정보 사업)",
        "위치정보 이용약관 별도 동의 (가입 시)",
        "위치 데이터 보관 기간 최소화 (인증 후 좌표 익명화)",
        "미성년자 이용 제한 (만 14세 미만 보호자 동의)",
    ]:
        add_bullet(doc, item)

    # 15. 리스크
    add_heading1(doc, "15. 리스크 및 대응 방안")
    add_hr(doc)

    make_table(doc,
        headers=["리스크", "내용", "대응 방안"],
        rows=[
            ["GPS 정확도", "실내·도심 위치 오류", "반경 조정(50~200m), 사진 보조 인증(V2+)"],
            ["사용자 어뷰징", "GPS 조작 허위 방문", "이동 패턴 이상 감지, 신고 기능"],
            ["배터리 소모", "GPS 상시 사용 시", "여행 모드 ON/OFF, 백그라운드 미사용"],
            ["콘텐츠 부족", "배지 부족으로 조기 이탈", "MVP 전 100곳 준비, UGC 가속"],
            ["경쟁 서비스 출현", "대기업·유사 서비스 진입", "배지·커뮤니티·B2B 차별화"],
            ["개인정보 민감성", "위치정보 수집 거부감", "여행 모드 명시적 ON/OFF, 정책 투명"],
            ["광고 도입 부작용", "광고로 사용자 이탈", "Phase 1.5 비침습적, 구독자 제거"],
            ["지도 SDK 변경", "API 정책·요금 변경", "이중화 설계, 추상화 레이어"],
            ["법적 리스크", "위치정보법·개인정보법", "출시 전 전문가 검토, 사업자 신고"],
        ],
        col_widths=[3.5, 5.5, 7],
    )

    doc.add_page_break()

    # 부록 A
    add_heading1(doc, "부록 A. MVP 개발 체크리스트")
    add_hr(doc)
    add_checkbox_list(doc, [
        "Voyna 도메인·상표권 등록 (voyna.app / voyna.co.kr / 키프리스)",
        "Apple Developer 계정 등록 ($99/년)",
        "Supabase 프로젝트 생성 및 DB 스키마 구현 (v2.0)",
        "FlutterFlow 프로젝트 생성 및 Supabase 연동",
        "카카오맵 API 키 발급 (MVP) + 네이버맵 사전 신청",
        "회원가입/로그인 (카카오 소셜 로그인) 구현",
        "초기 여행지 DB 50곳 입력 (서울 명소, 부록 C 참조)",
        "배지 디자인 50종 제작 (등급별 프레임 + 명소별)",
        "GPS 여행 모드 ON/OFF 기능 구현",
        "자동 배지 획득 로직 구현 (일반)",
        "확인형 배지 획득 로직 구현 (희귀·특별)",
        "홈·맵·배지·더보기 4탭 화면 구현",
        "XP 및 레벨 표시 구현",
        "테스트 (실외 GPS 현장 테스트)",
        "TestFlight(iOS) / 내부 테스트(Android) 배포",
        "베타 테스터 100명 피드백 수집",
        "개인정보 처리방침·위치정보 이용약관 게시",
        "App Store / Google Play 정식 제출",
    ])

    # 부록 B
    add_heading1(doc, "부록 B. 주요 지표 (KPI)")
    add_hr(doc)

    make_table(doc,
        headers=["단계", "목표 지표"],
        rows=[
            ["MVP 런칭 1개월", "가입자 500명, DAU 100명, 1인당 평균 배지 3개"],
            ["3개월", "MAU 3,000명, 리텐션 D30 20% 이상"],
            ["6개월", "MAU 10,000명, 광고 월 수익 50만 원~"],
            ["12개월", "MAU 30,000명, 구독 전환율 5%, MRR 500만 원"],
            ["24개월", "MAU 100,000명, MRR 3,000만 원, B2B 파트너 50곳"],
        ],
        col_widths=[4.5, 11.5],
    )
    add_quote(doc, "North Star Metric: 월간 배지 획득 수 (Monthly Badges Acquired)")

    doc.add_page_break()

    # 부록 C - 서울 50곳 (신규)
    add_heading1(doc, "부록 C. 서울 초기 명소 50곳 (요약)")
    add_hr(doc)
    add_body(doc, "MVP 런칭 시점 입력 대상. 자세한 내용은 별도 파일 `data/seoul_50_locations.csv` 참조.")

    add_heading2(doc, "등급 분포")
    make_table(doc,
        headers=["등급", "개수", "XP", "비고"],
        rows=[
            ["일반 (Common)", "38곳", "+50", "자동 획득"],
            ["희귀 (Rare)", "9곳", "+150", "'확인' 버튼 팝업"],
            ["특별 (Special)", "3곳", "+300~400", "메인 랜드마크, 화려한 연출"],
            ["**합계**", "**50**", "—", "—"],
        ],
        col_widths=[5, 3, 3, 5],
    )

    add_heading2(doc, "카테고리 분포")
    make_table(doc,
        headers=["카테고리", "개수", "주요 명소"],
        rows=[
            ["역사 (고궁·문화재)", "10", "경복궁⭐, 창덕궁💎, 종묘💎, 청와대⭐"],
            ["한옥/거리", "5", "북촌💎, 인사동, 익선동"],
            ["랜드마크", "7", "N서울타워⭐, 롯데월드타워💎, DDP"],
            ["시장", "3", "광장시장, 남대문, 망원시장"],
            ["거리/번화가", "5", "명동, 홍대, 강남역, 이태원"],
            ["한강 공원", "7", "여의도, 반포💎, 잠실, 노들섬💎"],
            ["자연/공원", "7", "북한산⭐, 인왕산💎, 청계산💎, 서울숲"],
            ["박물관/문화", "5", "국립중앙박물관💎, 리움💎, MMCA"],
            ["기타", "1", "청계광장"],
        ],
        col_widths=[4, 2, 10],
    )
    add_quote(doc, "⭐ 특별 등급 (3곳), 💎 희귀 등급 (9곳)")

    # 맺음말
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("이 기획서는 초기 기획 단계의 내부 문서이며, 시장 반응에 따라 지속적으로 업데이트됩니다.")
    set_font(run, size=9.5, italic=True, color=(120, 120, 120))

    closing2 = doc.add_paragraph()
    closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = closing2.add_run("v2.0 (2026-05-12) — Voyna 브랜드 리뉴얼, 화면 구성 명세화, 배지 획득 이원화, 경쟁사·개발 일정·개인정보 정책 추가")
    set_font(run2, size=8.5, italic=True, color=(150, 150, 150))

    doc.save(OUTPUT_PATH)
    print(f"DOCX saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
